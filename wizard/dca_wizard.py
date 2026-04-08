import base64
import io
import json
import logging
import os
import re
import unicodedata

from odoo import _, fields, models
from odoo.exceptions import UserError

try:
    from docx import Document
except ImportError:
    Document = None

_logger = logging.getLogger(__name__)


class DcaWizard(models.TransientModel) :

    _name = 'dca.wizard'

    report_models = fields.Selection([('alten', 'Alten'), ('simplified', 'Achmitech')], 'model de dossier',  required=True, default='simplified')
    applicant_id = fields.Many2one('hr.applicant', 'candidate')
    code_job = fields.Char('Reference de poste')
    report_type = fields.Selection([('PDF', 'pdf'), ('WORD', 'word')], 'type de dossier', required=True, default='WORD')
    _DCA_CATEGORY_ORDER = [
        'Soft Skills',
        'Langues',
        'Logiciels',
        'Langages de programmation',
        'Matériels',
        'Méthodes',
        'Normes et protocoles',
        'Systèmes',
        'Technologies',
        'Marketing',
    ]
    _LANGUAGE_LEVELS = {'A1', 'A2', 'B1', 'B2', 'C1', 'C2'}


    def get_applicant_extracted_payload(self) :
        raw_payload = self.applicant_id.applicant_extracted_json or '{}'
        if isinstance(raw_payload, dict) :
            return raw_payload
        
        try:
            payload = json.loads(raw_payload)
            return payload if isinstance(payload, dict) else {}
        except Exception:
            return {}

    # conversion robuste vers chaîne nettoyée.
    def _to_text(self, value):
        return str(value or '').strip()

    # normalise un label (minuscules, suppression accents/punctuation) pour matcher des titres Word.
    def _normalize_label(self, text):
        raw_text = self._to_text(text).lower()
        normalized = ''.join(
            ch for ch in unicodedata.normalize('NFD', raw_text)
            if unicodedata.category(ch) != 'Mn'
        )
        normalized = re.sub(r'[^a-z0-9]+', ' ', normalized)
        return re.sub(r'\s+', ' ', normalized).strip()

    # écrit proprement dans un paragraphe (gestion des runs).
    def _set_paragraph_text(self, paragraph, text):
        value = self._to_text(text)
        if paragraph.runs:
            paragraph.runs[0].text = value
            for run in paragraph.runs[1:]:
                run.text = ''
            return
        paragraph.text = value

    # écrit proprement dans une cellule de table.
    def _set_cell_text(self, cell, text):
        value = self._to_text(text)
        if cell.paragraphs:
            self._set_paragraph_text(cell.paragraphs[0], value)
            for paragraph in cell.paragraphs[1:]:
                self._set_paragraph_text(paragraph, '')
            return
        cell.text = value

    # détecte les lignes de points (.... ou ...) utilisées comme zones de saisie visuelle.
    def _is_dotted_placeholder(self, text):
        raw_text = self._to_text(text)
        if not raw_text:
            return False
        compact = re.sub(r'\s+', '', raw_text)
        return bool(re.fullmatch(r'[\.\u2026]{2,}', compact))

    #  remplit le bloc pointillé suivant un label narratif et retourne un booléen (True/False).
    def _fill_following_dotted_block(self, paragraphs, start_index, value):
        dotted_indexes = []
        label_markers = (
            'contexte general',
            'sujet du projet',
            'responsabilites occupees',
            'travail realise',
            'resultats obtenus',
        )

        for idx in range(start_index, len(paragraphs)):
            paragraph_text = self._to_text(paragraphs[idx].text)
            normalized = self._normalize_label(paragraph_text)
            if any(marker in normalized for marker in label_markers):
                break
            if self._is_dotted_placeholder(paragraph_text):
                dotted_indexes.append(idx)
                continue
            if dotted_indexes and paragraph_text:
                break

        if not dotted_indexes:
            return False

        raw_lines = [line.strip() for line in str(value or '').split('\n') if line.strip()]
        if not raw_lines:
            raw_lines = ['—']

        for pos, paragraph_idx in enumerate(dotted_indexes):
            if pos < len(raw_lines):
                self._set_paragraph_text(paragraphs[paragraph_idx], raw_lines[pos])
            else:
                self._set_paragraph_text(paragraphs[paragraph_idx], '')

        # Keep remaining content if list has more items than template lines.
        if len(raw_lines) > len(dotted_indexes):
            overflow = '\n'.join(raw_lines[len(dotted_indexes):])
            last_idx = dotted_indexes[-1]
            base_text = self._to_text(paragraphs[last_idx].text)
            self._set_paragraph_text(paragraphs[last_idx], '%s\n%s' % (base_text, overflow))
        return True

    #  fallback ajouté pour écrire directement sur la ligne du label si aucun bloc pointillé n’est détecté.
    def _fill_label_inline(self, paragraph, value):
        clean_value = self._to_text(value)
        if not clean_value:
            return

        current = self._to_text(paragraph.text)
        if not current:
            self._set_paragraph_text(paragraph, clean_value)
            return

        if ':' in current:
            label = current.split(':', 1)[0].strip()
            self._set_paragraph_text(paragraph, '%s: %s' % (label, clean_value))
            return

        self._set_paragraph_text(paragraph, '%s %s' % (current, clean_value))

    def _remove_table(self, table):
        table_element = table._element
        parent_element = table_element.getparent()
        if parent_element is not None:
            parent_element.remove(table_element)

    def _remove_paragraph(self, paragraph):
        paragraph_element = paragraph._element
        parent_element = paragraph_element.getparent()
        if parent_element is not None:
            parent_element.remove(paragraph_element)

    def _clear_paragraph_numbering(self, paragraph):
        paragraph_properties = paragraph._element.pPr
        if paragraph_properties is not None and paragraph_properties.numPr is not None:
            paragraph_properties.remove(paragraph_properties.numPr)

    def _clear_paragraph_page_breaks(self, paragraph):
        paragraph_properties = paragraph._element.pPr
        if paragraph_properties is not None and paragraph_properties.pageBreakBefore is not None:
            paragraph_properties.remove(paragraph_properties.pageBreakBefore)

        for run in paragraph.runs:
            run_element = run._element
            for break_element in list(run_element.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br')):
                run_element.remove(break_element)

    def _prune_unused_project_sections(self, document, experiences_count):
        project_start_table_idx = 3
        project_pair_step = 2

        max_experiences = max(0, int(experiences_count or 0))
        first_table_to_remove = project_start_table_idx + (max_experiences * project_pair_step)
        if first_table_to_remove < len(document.tables):
            for table in reversed(document.tables[first_table_to_remove:]):
                self._remove_table(table)

        paragraphs = document.paragraphs
        current_experience_idx = -1
        in_pruned_experience_block = False
        paragraphs_to_remove = []

        for idx, paragraph in enumerate(paragraphs):
            normalized = self._normalize_label(paragraph.text)
            if 'contexte general' in normalized:
                current_experience_idx += 1
                in_pruned_experience_block = current_experience_idx >= max_experiences

            if not in_pruned_experience_block:
                continue

            self._set_paragraph_text(paragraph, '')
            self._clear_paragraph_numbering(paragraph)
            self._clear_paragraph_page_breaks(paragraph)
            paragraphs_to_remove.append(paragraph)

            # Remove dotted placeholders immediately after pruned labels.
            if any(marker in normalized for marker in (
                'contexte general',
                'sujet du projet',
                'responsabilites occupees',
                'travail realise',
                'resultats obtenus',
            )):
                self._fill_following_dotted_block(paragraphs, idx + 1, '')

        # Physically remove pruned paragraphs so Word no longer reserves blank pages.
        for paragraph in reversed(paragraphs_to_remove):
            self._remove_paragraph(paragraph)

    # nettoie et déduplique les tâches
    def _normalize_task_list(self, tasks):
        if isinstance(tasks, str):
            return [self._to_text(tasks)] if self._to_text(tasks) else []
        if not isinstance(tasks, list):
            return []

        normalized = []
        seen = set()
        for raw_task in tasks:
            task_text = self._to_text(raw_task)
            if not task_text:
                continue
            fingerprint = task_text.lower()
            if fingerprint in seen:
                continue
            seen.add(fingerprint)
            normalized.append(task_text)
        return normalized

    # get the infos applicant.
    def _normalize_applicant_payload_for_word(self, applicant_data):
        payload = applicant_data if isinstance(applicant_data, dict) else {}

        education_raw = payload.get('education') if isinstance(payload.get('education'), dict) else {}
        education = {
            'degree': self._to_text(education_raw.get('degree')),
            'field': self._to_text(education_raw.get('field')),
            'university': self._to_text(education_raw.get('university')),
        }

        skills_raw = payload.get('skills') if isinstance(payload.get('skills'), dict) else {}
        skills = {}
        for raw_name, raw_level in skills_raw.items():
            name = self._to_text(raw_name)
            level = self._to_text(raw_level)
            if name and level:
                skills[name] = level

        certifications = []
        seen_certifications = set()
        for raw_certification in (payload.get('certification') or []):
            certification_text = self._to_text(raw_certification)
            if not certification_text:
                continue
            certification_key = certification_text.lower()
            if certification_key in seen_certifications:
                continue
            seen_certifications.add(certification_key)
            certifications.append(certification_text)

        normalized_experiences = []
        for raw_experience in (payload.get('experiences') or []):
            if not isinstance(raw_experience, dict):
                continue

            normalized_skills = {}
            raw_skills_per_category = raw_experience.get('skills_pertinents')
            if isinstance(raw_skills_per_category, dict):
                for category_name in self._DCA_CATEGORY_ORDER:
                    bucket = []
                    seen_bucket = set()
                    for raw_item in (raw_skills_per_category.get(category_name) or []):
                        skill_item = self._to_text(raw_item)
                        if not skill_item:
                            continue
                        skill_key = skill_item.lower()
                        if skill_key in seen_bucket:
                            continue
                        seen_bucket.add(skill_key)
                        bucket.append(skill_item)
                    normalized_skills[category_name] = bucket
            else:
                normalized_skills = {category_name: [] for category_name in self._DCA_CATEGORY_ORDER}

            tasks = self._normalize_task_list(raw_experience.get('tasks'))

            normalized_experiences.append({
                'company': self._to_text(raw_experience.get('company')),
                'duration': self._to_text(raw_experience.get('duration')),
                'general_context': self._pick_experience_field_text(raw_experience, 'general_context'),
                'project_topic': self._pick_experience_field_text(raw_experience, 'project_topic'),
                'responsibilities': self._pick_experience_field_text(raw_experience, 'responsibilities'),
                'work_done': self._pick_experience_field_text(raw_experience, 'work_done'),
                'results_obtained': self._pick_experience_field_text(raw_experience, 'results_obtained'),
                'tasks': tasks,
                'skills_pertinents': normalized_skills,
            })

        return {
            'name': self._to_text(payload.get('name')),
            'education': education,
            'skills': skills,
            'experiences': normalized_experiences,
            'certification': certifications,
            'experience_years': self._to_text(payload.get('experience_years')),
        }

    # récupère un champ narratif avec aliases (general_context, contexte_general, etc.). Si vide, construit une phrase à partir des tasks.
    def _pick_experience_field_text(self, experience, field_name):
        if not isinstance(experience, dict):
            return ''

        def _list_to_lines(items):
            cleaned_items = [self._to_text(item) for item in (items or []) if self._to_text(item)]
            return '\n'.join(cleaned_items)

        alias_map = {
            'general_context': (
                'general_context', 'contexte_general', 'generalContext', 'context', 'contexte',
            ),
            'project_topic': (
                'project_topic', 'sujet_du_projet', 'projectTopic', 'project_subject', 'topic',
            ),
            'responsibilities': (
                'responsibilities', 'responsabilites_occupees', 'responsibility', 'role', 'missions',
            ),
            'work_done': (
                'work_done', 'travail_realise', 'workDone', 'realisations',
            ),
            'results_obtained': (
                'results_obtained', 'resultats_obtenus', 'resultsObtained', 'results',
            ),
        }

        for key_name in alias_map.get(field_name, (field_name,)):
            raw_value = experience.get(key_name)
            if isinstance(raw_value, list):
                if field_name in ('work_done', 'results_obtained'):
                    value = _list_to_lines(raw_value)
                else:
                    value = '; '.join(self._to_text(item) for item in raw_value if self._to_text(item))
            else:
                value = self._to_text(raw_value)
            if value:
                return value

        tasks = self._normalize_task_list(experience.get('tasks'))
        if not tasks:
            return ''

        tasks_block = '; '.join(tasks)
        if field_name == 'general_context':
            return 'Contexte de mission deduit des taches realisees: %s.' % tasks_block
        if field_name == 'project_topic':
            return 'Sujet principal du projet deduit des taches realisees: %s.' % tasks_block
        if field_name == 'responsibilities':
            return 'Responsabilites occupees deduites des taches realisees: %s.' % tasks_block
        if field_name == 'work_done':
            return _list_to_lines(tasks)
        if field_name == 'results_obtained':
            return _list_to_lines(tasks)

        return ''

    # remplit le docx (alten/simplified) par structure.
    def _fill_dca_template_by_layout(self, document, applicant_data):
        skills = applicant_data.get('skills') if isinstance(applicant_data.get('skills'), dict) else {}
        education = applicant_data.get('education') if isinstance(applicant_data.get('education'), dict) else {}
        experiences = applicant_data.get('experiences') if isinstance(applicant_data.get('experiences'), list) else []
        certifications = [self._to_text(item) for item in (applicant_data.get('certification') or []) if self._to_text(item)]

        category_values = {}
        for category_name in self._DCA_CATEGORY_ORDER:
            if category_name == 'Langues':
                category_values[category_name] = [
                    self._to_text(skill_name)
                    for skill_name, level in skills.items()
                    if self._to_text(level).upper() in self._LANGUAGE_LEVELS and self._to_text(skill_name)
                ]
                continue

            category_bucket = []
            for experience in experiences:
                if not isinstance(experience, dict):
                    continue
                category_dict = experience.get('skills_pertinents') if isinstance(experience.get('skills_pertinents'), dict) else {}
                for item in (category_dict.get(category_name) or []):
                    cleaned = self._to_text(item)
                    if cleaned and cleaned not in category_bucket:
                        category_bucket.append(cleaned)
            category_values[category_name] = category_bucket

        if len(document.tables) > 1:
            competencies_table = document.tables[1]
            table_to_category = {
                'logiciels': 'Logiciels',
                'langages': 'Langages de programmation',
                'technologies': 'Technologies',
                'methodes': 'Méthodes',
                'systemes': 'Systèmes',
                'normes et protocoles': 'Normes et protocoles',
                'materiels': 'Matériels',
            }
            for row in competencies_table.rows:
                if len(row.cells) < 2:
                    continue
                label = self._normalize_label(row.cells[0].text)
                category_name = table_to_category.get(label)
                if not category_name:
                    continue
                values = category_values.get(category_name) or []
                self._set_cell_text(row.cells[1], ', '.join(values) if values else '—')

        if len(document.tables) > 2:
            formation_table = document.tables[2]
            education_parts = [
                self._to_text(education.get('degree')),
                self._to_text(education.get('field')),
                self._to_text(education.get('university')),
            ]
            education_parts = [part for part in education_parts if part]
            language_values = category_values.get('Langues') or []
            for row in formation_table.rows:
                if len(row.cells) < 2:
                    continue
                row_label = self._normalize_label(row.cells[0].text)
                if row_label == 'formation':
                    self._set_cell_text(row.cells[1], ' - '.join(education_parts) if education_parts else '—')
                elif row_label == 'habilitations':
                    self._set_cell_text(row.cells[1], ', '.join(certifications) if certifications else '—')
                elif row_label == 'langues':
                    self._set_cell_text(row.cells[1], ', '.join(language_values) if language_values else '—')

        self._prune_unused_project_sections(document, len(experiences))

        project_start_table_idx = 3
        project_pair_step = 2
        project_slots = max(0, (len(document.tables) - project_start_table_idx) // project_pair_step)

        for project_idx in range(min(project_slots, len(experiences))):
            experience = experiences[project_idx] if project_idx < len(experiences) and isinstance(experiences[project_idx], dict) else {}

            header_table_idx = project_start_table_idx + (project_idx * project_pair_step)
            env_table_idx = header_table_idx + 1

            if header_table_idx < len(document.tables):
                header_table = document.tables[header_table_idx]
                if header_table.rows and len(header_table.rows[0].cells) >= 2:
                    self._set_cell_text(header_table.rows[0].cells[0], self._to_text(experience.get('company')) or '—')
                    self._set_cell_text(header_table.rows[0].cells[1], self._to_text(experience.get('duration')) or '—')

            if env_table_idx < len(document.tables):
                env_table = document.tables[env_table_idx]
                row_map = {
                    'logiciels': 'Logiciels',
                    'langages': 'Langages de programmation',
                    'technologies': 'Technologies',
                    'methodes': 'Méthodes',
                    'systemes': 'Systèmes',
                    'normes et protocoles': 'Normes et protocoles',
                    'materiels': 'Matériels',
                }
                skills_pertinents = experience.get('skills_pertinents') if isinstance(experience.get('skills_pertinents'), dict) else {}
                for row in env_table.rows:
                    if len(row.cells) < 3:
                        continue
                    row_label = self._normalize_label(row.cells[1].text)
                    category_name = row_map.get(row_label)
                    if not category_name:
                        continue
                    values = [
                        self._to_text(item)
                        for item in (skills_pertinents.get(category_name) or [])
                        if self._to_text(item)
                    ]
                    self._set_cell_text(row.cells[2], ', '.join(values) if values else '—')

        paragraphs = document.paragraphs
        current_experience_idx = -1
        label_to_field = {
            'contexte general': 'general_context',
            'sujet du projet': 'project_topic',
            'responsabilites occupees': 'responsibilities',
            'travail realise': 'work_done',
            'resultats obtenus': 'results_obtained',
        }

        for idx, paragraph in enumerate(paragraphs):
            normalized = self._normalize_label(paragraph.text)
            if 'contexte general' in normalized:
                current_experience_idx += 1

            field_name = None
            for label, mapped_field in label_to_field.items():
                if label in normalized:
                    field_name = mapped_field
                    break

            if field_name is None or current_experience_idx < 0:
                continue

            if current_experience_idx >= len(experiences):
                self._fill_following_dotted_block(paragraphs, idx + 1, '')
                continue

            experience = experiences[current_experience_idx] if isinstance(experiences[current_experience_idx], dict) else {}
            resolved_value = self._pick_experience_field_text(experience, field_name) or '—'
            filled = self._fill_following_dotted_block(
                paragraphs,
                idx + 1,
                resolved_value,
            )
            if not filled:
                self._fill_label_inline(paragraph, resolved_value)

    # résout le chemin du template selon report_models.
    def _get_template_path(self):
        module_root = os.path.dirname(os.path.dirname(__file__))
        template_by_model = {
            'alten': 'dca_alten_template.docx',
            'simplified': 'dca_simplified_template.docx',
        }
        template_name = template_by_model.get(self.report_models)
        if not template_name:
            raise UserError(_('Unsupported report model for Word export.'))

        template_path = os.path.join(module_root, 'static', 'src', 'docx', template_name)
        if not os.path.exists(template_path):
            raise UserError(_('Word template not found: %s') % template_name)
        return template_path

    # code job add function.
    def _fill_alten_footer_code_job(self, document, code_job_value):
        footer_value = self._to_text(code_job_value)
        if not footer_value:
            return

        for section in document.sections:
            for table in section.footer.tables:
                for row in table.rows:
                    if len(row.cells) < 1:
                        continue
                    self._set_cell_text(row.cells[0], footer_value)
                    return

    # pipeline.
    def _render_word_report(self, applicant_data, code_job_value=''):
        if Document is None:
            raise UserError(_('python-docx is not installed. Please install dependency: docx'))

        normalized_applicant_data = self._normalize_applicant_payload_for_word(applicant_data)
        template_path = self._get_template_path()

        document = Document(template_path)
        if self.report_models in ('alten', 'simplified'):
            self._fill_dca_template_by_layout(document, normalized_applicant_data)
        if self.report_models == 'alten':
            self._fill_alten_footer_code_job(document, code_job_value)

        output_stream = io.BytesIO()
        document.save(output_stream)
        output_stream.seek(0)

        filename = 'Dossier de Compétences %s.docx' % (
            re.sub(r'[^A-Za-z0-9_\-]+', '_', self._to_text(self.applicant_id.partner_name) or 'candidate'),
        )

        attachment = self.env['ir.attachment'].create({
            'name': filename,
            'type': 'binary',
            'datas': base64.b64encode(output_stream.read()),
            'mimetype': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'res_model': self._name,
            'res_id': self.id,
        })
        applicant = self.env['hr.applicant'].search([('id' ,'=', self.applicant_id)])
        if applicant:
            applicant.message_post(attachment_ids=attachment.ids)

        return {
            'type': 'ir.actions.act_url',
            'url': '/web/content/%s?download=true' % attachment.id,
            'target': 'self',
        }

    def _get_or_extract_applicant_payload(self):
        applicant_data = self.get_applicant_extracted_payload()
        if applicant_data:
            return applicant_data

        applicant_data = self.applicant_id.get_extracted_applicant_data()
        if applicant_data:
            self.applicant_id.write({
                'applicant_extracted_json': json.dumps(applicant_data or {}, ensure_ascii=False, indent=2),
            })
        return applicant_data if isinstance(applicant_data, dict) else {}
    
    def print_report(self) :
        self.ensure_one()
        if not self.applicant_id:
            raise UserError(_('Please select a candidate before generating the report.'))

        applicant_data = self._get_or_extract_applicant_payload()
        code_job_value = (self.code_job or '').strip()

        if self.report_models == 'alten' and self.report_type == 'PDF' : 
            return self.env.ref('scoring_candidates.action_print_dca_alten').with_context(
                dca_code_job=code_job_value
            ).report_action(self.applicant_id)
        
        elif self.report_models == 'simplified' and self.report_type == 'PDF':
            return self.env.ref('scoring_candidates.action_print_dca_simplify').with_context(
                dca_code_job=code_job_value
            ).report_action(self.applicant_id)
        
        if self.report_models == 'alten' and self.report_type == 'WORD' : 
            return self._render_word_report(applicant_data, code_job_value)
        
        elif self.report_models == 'simplified' and self.report_type == 'WORD':
            return self._render_word_report(applicant_data, code_job_value)
        
        return True