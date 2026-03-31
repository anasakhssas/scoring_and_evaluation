import io
import logging
import os
from docx import Document

_logger = logging.getLogger(__name__)


def generate_dossier_docx(json_data):
    """Generate dossier de competences docx from extracted applicant JSON."""
    doc = Document()
    doc.add_heading('DOSSIER DE COMPETENCES', level=1)

    # candidate core info
    table0 = doc.add_table(rows=4, cols=2)
    table0.style = 'Table Grid'
    table0.cell(0, 0).text = 'Nom Prénom'
    table0.cell(1, 0).text = 'Dernier Diplôme'
    table0.cell(2, 0).text = 'Années d\'Expériences'
    table0.cell(3, 0).text = 'Date de Disponibilité'

    table0.cell(0, 1).text = str(json_data.get('name') or '')
    education = json_data.get('education') or {}
    diploma = ' '.join([str(education.get(key) or '') for key in ('degree', 'field', 'university')]).strip()
    table0.cell(1, 1).text = diploma
    table0.cell(2, 1).text = str(json_data.get('experience_years') or '')
    table0.cell(3, 1).text = str(json_data.get('availability') or '')

    skills = json_data.get('skills') or {}
    items = list(skills.items()) if isinstance(skills, dict) else []
    lang_list = [it for it in items if it[0].lower() in ('english', 'french', 'spanish', 'arabic', 'german', 'italian', 'portuguese', 'chinese', 'japanese', 'korean', 'russian', 'dutch', 'turkish')]
    other_list = [it for it in items if it not in lang_list]

    def add_skill_section(title, skill_list):
        doc.add_paragraph(title, style='Heading 2')
        tbl = doc.add_table(rows=len(skill_list)+1, cols=2)
        tbl.style = 'Table Grid'
        tbl.cell(0, 0).text = 'Compétence'
        tbl.cell(0, 1).text = 'Niveau'
        for i, (name, level) in enumerate(skill_list, start=1):
            tbl.cell(i, 0).text = str(name)
            tbl.cell(i, 1).text = str(level)

    add_skill_section('VOS TOP 5 DES HARD SKILLS', other_list[0:5])
    add_skill_section('VOS TOP 5 DES SOFT SKILLS', other_list[5:10])
    add_skill_section('VOS TOP 5 DES OUTILS MAITRISES', other_list[10:15])
    add_skill_section('VOS LANGUES MAITRISES', lang_list[0:5])

    experiences = json_data.get('experiences') or []
    if experiences:
        doc.add_heading('EXPERIENCES', level=2)
        for exp in experiences[:5]:
            p = doc.add_paragraph(style='List Number')
            p.add_run('Poste: ').bold = True
            p.add_run(str(exp.get('title') or ''))
            doc.add_paragraph('Entreprise: %s' % str(exp.get('company') or ''))
            doc.add_paragraph('Durée: %s' % str(exp.get('duration') or ''))
            if exp.get('tasks'):
                doc.add_paragraph('Mission:')
                for task in exp.get('tasks'):
                    doc.add_paragraph(str(task), style='List Bullet')
            if exp.get('achievements'):
                doc.add_paragraph('Réalisations:')
                for ach in exp.get('achievements'):
                    doc.add_paragraph(str(ach), style='List Bullet')

    doc.add_heading('FORMATIONS', level=2)
    if education:
        doc.add_paragraph('%s %s %s' % (education.get('degree', ''), education.get('field', ''), education.get('university', '')))
    else:
        doc.add_paragraph('N/A')

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()

    skills = json_data.get('skills') or {}
    items = list(skills.items()) if isinstance(skills, dict) else []

    lang_list = [it for it in items if it[0].lower() in ('english', 'french', 'spanish', 'arabic', 'german', 'italian', 'portuguese', 'chinese', 'japanese', 'korean', 'russian', 'dutch', 'turkish')]
    other_list = [it for it in items if it not in lang_list]

    if len(doc.tables) >= 5:
        _fill_two_column_skill_table(doc.tables[1], other_list[0:5])
        _fill_two_column_skill_table(doc.tables[2], other_list[5:10])
        _fill_two_column_skill_table(doc.tables[3], other_list[10:15])
        _fill_two_column_skill_table(doc.tables[4], lang_list[0:5])

    experiences = json_data.get('experiences') or []
    for idx in range(5, min(10, 5 + len(experiences))):
        exp = experiences[idx - 5] if idx - 5 < len(experiences) else None
        if exp and idx < len(doc.tables):
            _fill_experience_table(doc.tables[idx], exp)

    if len(doc.tables) > 10:
        table10 = doc.tables[10]
        if len(table10.rows) > 1:
            _set_cell_text(table10.cell(1, 0), str(education.get('year') or ''))
            _set_cell_text(table10.cell(1, 1), diploma)

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()


def _set_cell_text(cell, text):
    cell.text = str(text or '')


def _fill_two_column_skill_table(table, skills):
    for row_idx, (name, level) in enumerate(skills, start=1):
        if row_idx >= len(table.rows):
            break
        _set_cell_text(table.cell(row_idx, 0), name)
        _set_cell_text(table.cell(row_idx, 1), level)


def _fill_experience_table(table, exp):
    start, end = '', ''
    duration = str(exp.get('duration') or '')
    if '-' in duration:
        duration_split = duration.split('-', 1)
        start, end = duration_split[0].strip(), duration_split[1].strip()

    _set_cell_text(table.cell(0, 0), f'DATE DEBUT : {start}')
    _set_cell_text(table.cell(0, 1), f"POSTE OCCUPE : {exp.get('title', '')}\nENTREPRISE : {exp.get('company', '')}")
    _set_cell_text(table.cell(0, 2), f'DATE FIN : {end}')

    mission_text = '\n'.join(f"- {task}" for task in (exp.get('tasks') or []))
    if len(table.rows) > 3:
        _set_cell_text(table.cell(3, 0), f"MISSION :\n{mission_text}")

    realization_text = '\n'.join(f"- {item}" for item in (exp.get('achievements') or []))
    if len(table.rows) > 5:
        _set_cell_text(table.cell(5, 0), f"REALISATIONS :\n{realization_text}")
